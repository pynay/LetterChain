from fastapi import UploadFile, HTTPException, status
import PyPDF2
import io
import hashlib
from typing import Tuple, Optional
import logging
from app.core.config import settings
from app.models.schemas import FileUploadResponse

logger = logging.getLogger(__name__)

class FileService:
    """Centralized file handling with validation and processing"""
    
    @staticmethod
    def validate_upload(
        file: UploadFile,
        file_bytes: bytes,
        file_label: str
    ) -> FileUploadResponse:
        """Validate file size, extension, and MIME type"""
        
        # Check file size
        if len(file_bytes) > settings.MAX_FILE_SIZE:
            return FileUploadResponse(
                filename=file.filename,
                size=len(file_bytes),
                content_type=file.content_type,
                is_valid=False,
                error_message=f"{file_label} file is too large (max {settings.MAX_FILE_SIZE // 1024 // 1024}MB)."
            )
        
        # Check file extension
        if "." in file.filename:
            ext = file.filename.lower().rsplit(".", 1)[-1]
            if f".{ext}" not in settings.ALLOWED_EXTENSIONS:
                return FileUploadResponse(
                    filename=file.filename,
                    size=len(file_bytes),
                    content_type=file.content_type,
                    is_valid=False,
                    error_message=f"{file_label} file type not allowed. Only {', '.join(settings.ALLOWED_EXTENSIONS)} are supported."
                )
        else:
            return FileUploadResponse(
                filename=file.filename,
                size=len(file_bytes),
                content_type=file.content_type,
                is_valid=False,
                error_message=f"{file_label} file must have a valid extension."
            )
        
        # Check MIME type
        if file.content_type not in settings.ALLOWED_MIME_TYPES:
            return FileUploadResponse(
                filename=file.filename,
                size=len(file_bytes),
                content_type=file.content_type,
                is_valid=False,
                error_message=f"{file_label} file MIME type not allowed."
            )
        
        return FileUploadResponse(
            filename=file.filename,
            size=len(file_bytes),
            content_type=file.content_type,
            is_valid=True
        )
    
    @staticmethod
    def extract_text(file: UploadFile, file_bytes: bytes) -> str:
        """Extract text content from uploaded file"""
        try:
            if file.filename.lower().endswith(".pdf"):
                return FileService._extract_pdf_text(file_bytes)
            elif file.filename.lower().endswith(".docx"):
                return FileService._extract_docx_text(file_bytes)
            else:
                return file_bytes.decode(errors="ignore")
        except Exception as e:
            logger.error(f"Failed to extract text from {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Error reading file {file.filename}: {str(e)}"
            )
    
    @staticmethod
    def _extract_pdf_text(file_bytes: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
                    continue
            
            if not text_parts:
                raise ValueError("No text could be extracted from PDF")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    @staticmethod
    def _extract_docx_text(file_bytes: bytes) -> str:
        """Extract text from Word document (.docx) file"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_bytes))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                raise ValueError("No text could be extracted from Word document")
            
            return "\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx library not installed")
            raise ValueError("Word document processing not available - python-docx library required")
        except Exception as e:
            logger.error(f"Word document processing failed: {str(e)}")
            raise ValueError(f"Failed to process Word document: {str(e)}")
    
    @staticmethod
    def generate_file_hash(content: str) -> str:
        """Generate SHA-256 hash of file content for caching"""
        return hashlib.sha256(content.encode()).hexdigest()
    
    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filename for safe file operations"""
        import re
        # Remove or replace problematic characters
        sanitized = re.sub(r'[^\w\-_.]', '_', filename)
        # Limit length
        if len(sanitized) > 100:
            name, ext = sanitized.rsplit('.', 1) if '.' in sanitized else (sanitized, '')
            sanitized = name[:95] + '.' + ext
        return sanitized

# Global file service instance
file_service = FileService()